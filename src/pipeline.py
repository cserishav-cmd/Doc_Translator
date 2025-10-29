import os
import fitz  # Import PyMuPDF
import re    # Import regular expression module
import uuid  # Import UUID module
from src.translator import translate_elements
from src.document_analyzer import DocumentAnalyzer  # Add document analyzer
# UPDATED: Import the new in-place PDF rebuilder
from src.rebuild import rebuild_pdf_in_place, rebuild_docx_with_lxml, convert_pdf_to_docx, convert_docx_to_pdf
# --- FIX: Import the Document class type specifically ---
from docx import Document as DocxDocument # Use alias to avoid conflict if needed later
from docx.document import Document as DocumentClass # Import the class type itself
# NEW: Import for lxml text box parsing
from docx.oxml.ns import qn
from collections import Counter
import traceback # For detailed error logging

# OCR disabled for speed optimization - using fast API-based translation only
# To re-enable OCR, uncomment the functions below and imports at top of file

# ---------------- Helper functions ----------------

def extract_pdf_elements(file_path):
    """
    ENHANCED: Extracts text from PDF using PyMuPDF with OCR fallback for scanned pages.
    Preserves images, coordinates, and comprehensive style information.
    """
    doc = None # Initialize doc to None
    elements = []
    images_info = []  # Track image locations to avoid overwriting

    try:
        doc = fitz.open(file_path)
        print(f"--- Starting PDF Extraction for: {os.path.basename(file_path)} ---") # Added logging
        # PyMuPDF's Document is iterable - linter doesn't recognize this
        for page_num, page in enumerate(doc):  # type: ignore[arg-type]
            print(f"  Processing PDF Page {page_num + 1}/{doc.page_count}") # Added logging
            # --- 1. Extract and catalog all images on this page ---
            image_list = page.get_images(full=True)
            page_image_count = 0 # Added logging
            for img_index, img in enumerate(image_list):
                xref = img[0]
                try:
                    # get_image_bbox can return multiple rects, take the first one
                    bboxes = page.get_image_bbox(img)
                    if bboxes:
                        # Ensure bbox is a tuple (x0, y0, x1, y1)
                        bbox_tuple = bboxes if isinstance(bboxes, tuple) else bboxes[0]
                        images_info.append({
                            "page_num": page_num,
                            "bbox": bbox_tuple,
                            "xref": xref
                        })
                        page_image_count += 1 # Added logging
                except Exception as img_e:
                    print(f"⚠️ Warning: Could not get bbox for image {xref} on page {page_num+1}: {img_e}")
            # if page_image_count > 0: print(f"    Found {page_image_count} images on page.") # Added logging


            # --- 2. Attempt digital text extraction ---
            page_dict = page.get_text("dict", flags=fitz.TEXTFLAGS_DICT & ~fitz.TEXT_PRESERVE_LIGATURES & ~fitz.TEXT_PRESERVE_WHITESPACE) # Flags to potentially improve spacing
            text_blocks = []
            has_digital_text = False
            page_block_count = 0 # Added logging

            for block in page_dict.get("blocks", []):
                if block["type"] == 0:  # Text block
                    block_text = ""
                    font_sizes = []
                    font_names = []
                    font_flags = []  # To detect bold/italic
                    span_texts = [] # Store individual span texts to handle spacing better

                    for line in block.get("lines", []):
                        line_text = ""
                        for span in line.get("spans", []):
                            span_text = span.get("text", "")
                            # Basic whitespace normalization within spans
                            normalized_span_text = ' '.join(span_text.split())
                            if normalized_span_text: # Only add if not just whitespace
                                line_text += normalized_span_text + " " # Add space between spans
                                font_sizes.append(round(span["size"]))
                                font_names.append(span.get("font", "helv"))
                                font_flags.append(span.get("flags", 0))
                        # Add line text to block text if it's not empty, add newline
                        stripped_line_text = line_text.strip()
                        if stripped_line_text:
                            block_text += stripped_line_text + "\n" # Use newline between lines

                    # Final cleanup for the block text
                    final_block_text = block_text.strip() # Remove leading/trailing whitespace including newlines

                    if final_block_text:
                        has_digital_text = True
                        page_block_count += 1 # Added logging
                        # Determine dominant font characteristics
                        common_size = Counter(font_sizes).most_common(1)[0][0] if font_sizes else 10
                        common_font = Counter(font_names).most_common(1)[0][0] if font_names else "helv"

                        # Detect bold/italic from flags (bit masking)
                        # Check flags more robustly
                        is_bold = any((f & (1 << 4)) for f in font_flags) # Bit 4 = bold
                        is_italic = any((f & (1 << 0)) for f in font_flags) # Bit 0 = italic

                        # Simple heuristic for headings (larger font size compared to common, maybe bold)
                        # This might need refinement based on document structure
                        is_heading_heuristic = (font_sizes and common_size > 14) or (is_bold and common_size > 12)


                        text_blocks.append({
                            "type": "paragraph", # Treat each block as paragraph for now
                            "text": final_block_text,
                            "is_heading": is_heading_heuristic, # Use heuristic
                            "page_num": page_num,
                            "bbox": block["bbox"],
                            "font_size": common_size,
                            "font_name": common_font,
                            "is_bold": is_bold,
                            "is_italic": is_italic,
                            "is_image_text": False # Assuming not from OCR here
                        })
            # print(f"    Found {page_block_count} text blocks on page.") # Added logging

            # --- 3. SKIP OCR (as before) ---
            if not has_digital_text:
                print(f"⚠️ Page {page_num+1} has minimal/no digital text (likely scanned/image).")
                print(f"   💡 TIP: For scanned PDFs, use online OCR (e.g., Adobe) before uploading.")
                # Optionally add placeholder
                # text_blocks.append({ ... placeholder dict ... })

            elements.extend(text_blocks)

        print(f"--- Finished PDF Extraction ---") # Added logging
        print(f"  Total elements extracted: {len(elements)}") # Added logging
        if images_info:
            print(f"  Total images found: {len(images_info)}") # Added logging

        # Remove trailing [EMPTY] placeholders if any were added and are last
        while elements and elements[-1].get("text", "") == "[EMPTY]":
            elements.pop()

        # <<< ADDED LOGGING >>>
        print(f"  Returning {len(elements)} elements from PDF extraction.")
        return elements

    except Exception as e:
        print(f"❌ Error extracting PDF elements from {file_path}: {e}")
        traceback.print_exc()
        return [] # Return empty list on error
    finally:
        if doc:
            try:
                doc.close()
            except Exception: pass # Ignore close errors

def extract_docx_elements_and_objects(file_path):
    """
    ENHANCED: Extracts text preserving paragraph structure more strictly.
    Adds logging to verify paragraph separation.
    """
    doc = None
    elements_to_translate = []
    standard_paragraph_objects = [] # List of high-level Paragraph objects
    textbox_paragraph_elements = [] # List of low-level lxml <w:p> elements

    print(f"\n--- Starting DOCX Extraction for: {os.path.basename(file_path)} ---")

    try:
        # Use the alias DocxDocument to open the file
        doc = DocxDocument(file_path)

        def process_text(text, is_heading, para_obj=None, lxml_obj=None, location=""):
            """Helper to create the translatable element dict."""
            stripped_text = text.strip()
            # If after stripping, the text is empty, represent it as [EMPTY]
            element_text = stripped_text if stripped_text else "[EMPTY]"
            # --- NEW: Detect if element is likely a mathematical equation ---
            def _looks_like_equation(s: str) -> bool:
                if not s or s == "[EMPTY]":
                    return False
                # Common LaTeX/math delimiters
                if re.search(r"(\\begin\{|\$\$|\$|\\\\\[|\\\\\(|\\\[|\\\()", s):
                    return True
                # Presence of many math symbols or operators suggests equation
                math_symbols = re.findall(r"[=<>\^_\\+\-/*%|:;{}\[\]\(\)]", s)
                # If lots of non-alpha characters relative to length -> equation
                ratio = len(math_symbols) / max(len(s), 1)
                if ratio > 0.12:
                    return True
                # Short single-line strings with digits and symbols often math
                tokens = s.split()
                digit_tokens = sum(1 for t in tokens if re.search(r"\d", t))
                if digit_tokens >= max(1, len(tokens) // 3) and len(tokens) < 20:
                    return True
                return False

            is_equation = _looks_like_equation(element_text)

            element = {
                "type": "paragraph",
                "text": element_text,
                "is_heading": is_heading,
                "is_equation": is_equation,
                "location": location # Added for logging
            }
            # --- Logging added ---
            # print(f"  Extracted ({location}): Head={is_heading}, Text='{element_text[:50].replace(chr(10), ' ')}...'") # Commented out for less noise
            return element

        def parse_paragraph_runs(para):
            """Helper to get text from a paragraph, preserving bold markers."""
            # Use para.text which often handles complex fields better, then strip
            # But fall back to run concatenation if para.text is empty but runs exist
            full_text_simple = para.text.strip()
            if full_text_simple:
                 # Reconstruct with bold markers if simple text exists
                 current_text = ""
                 for run in para.runs:
                      # Check bold status more reliably
                      is_bold = run.bold or (hasattr(run.font, 'bold') and run.font.bold is True)
                      run_text_content = run.text
                      # Handle potential None text in runs
                      if run_text_content is None: run_text_content = ""

                      if is_bold:
                          current_text += f"**{run_text_content}**"
                      else:
                          current_text += run_text_content
                 # Return reconstructed text if it has content, otherwise indicate empty
                 return current_text.strip() if current_text.strip() else "[EMPTY]"
            else:
                 # Check if runs exist even if para.text is empty
                 if para.runs:
                      current_text = ""
                      for run in para.runs:
                           is_bold = run.bold or (hasattr(run.font, 'bold') and run.font.bold is True)
                           run_text_content = run.text
                           if run_text_content is None: run_text_content = ""

                           if is_bold:
                               current_text += f"**{run_text_content}**"
                           else:
                               current_text += run_text_content
                      return current_text.strip() if current_text.strip() else "[EMPTY]"
                 else:
                      # Truly empty paragraph
                      return "[EMPTY]"


        def parse_lxml_paragraph(p_element):
            """
            Helper to get text from a low-level lxml <w:p> element, preserving bold.
            """
            text = ""
            for r_element in p_element.findall(qn('w:r')):
                run_text = ""
                is_bold = False
                rPr = r_element.find(qn('w:rPr'))
                if rPr is not None:
                    # Check for explicit bold tag <w:b/> or <w:b w:val="true"/>
                    b_tag = rPr.find(qn('w:b'))
                    if b_tag is not None:
                        # Check the 'val' attribute, default to True if tag exists
                        val = b_tag.get(qn('w:val'))
                        if val is None or val.lower() in ('1', 'true', 'on'):
                            is_bold = True

                # Concatenate text from all <w:t> elements within the run
                for t_element in r_element.findall(qn('w:t')):
                    if t_element.text:
                        run_text += t_element.text

                if is_bold:
                    text += f"**{run_text}**"
                else:
                    text += run_text

            return text.strip() if text.strip() else "[EMPTY]"

        # --- 1. Process Standard Content (Body, Tables, Headers, Footers) ---
        print("\n  Processing Standard Content...")
        all_standard_paras_objs = [] # Store objects to process
        # Body paragraphs
        all_standard_paras_objs.extend(doc.paragraphs)
        # Table paragraphs
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Process each paragraph within the cell distinctly
                    all_standard_paras_objs.extend(cell.paragraphs) # Append paragraph objects
        # Header/Footer paragraphs
        for section_idx, section in enumerate(doc.sections):
            try: # Add try-except for header/footer access
                if section.header:
                     all_standard_paras_objs.extend(section.header.paragraphs)
                if section.footer:
                     all_standard_paras_objs.extend(section.footer.paragraphs)
            except Exception as hf_e:
                print(f"  Warning: Could not access header/footer for section {section_idx}: {hf_e}")


        processed_para_ids = set() # Track processed paragraphs by object ID

        for idx, para in enumerate(all_standard_paras_objs):
             para_id = id(para)
             if para_id in processed_para_ids:
                 # print(f"  Skipping duplicate paragraph object reference at index {idx}.") # Optional: for debugging duplicates
                 continue # Skip if already processed

             # Determine if heading based on style name
             is_heading = para.style is not None and para.style.name.lower().startswith('heading')

             text = parse_paragraph_runs(para)

             # Add element for this paragraph ONLY
             location = f"Std Para {idx}" # Location based on unified list index
             print(f"  Extracting ({location}): Head={is_heading}, Text='{text[:50].replace(chr(10), ' ')}...'") # LOGGING
             elements_to_translate.append(process_text(text, is_heading, para_obj=para, location=location))
             standard_paragraph_objects.append(para)
             textbox_paragraph_elements.append(None) # Placeholder for alignment
             processed_para_ids.add(para_id) # Mark as processed


        print(f"  Finished Standard Content. Paragraphs processed: {len(standard_paragraph_objects)}")


        # --- 2. Process Text Box Content (lxml) ---
        print("\n  Processing Text Box Content...")
        textbox_paras_count = 0
        try:
            # Find all text box content containers using a more robust XPath
            # This looks for text boxes within shapes (w:txbx//w:p) and inline text boxes (w:drawing//w:txbxContent//w:p)
            # Adjust namespaces if necessary for your specific DOCX variant
            xpath_query = './/w:txbxContent//w:p | .//wps:txbx//w:p' # Example, might need refinement
            # Using findall directly on body might be simpler if the above is too complex or slow
            lxml_paragraphs_in_boxes = doc.element.body.findall(f'.//{qn("w:txbxContent")}//{qn("w:p")}')

            print(f"  Found {len(lxml_paragraphs_in_boxes)} paragraph elements within text boxes (using lxml).")

            for p_idx, p_element in enumerate(lxml_paragraphs_in_boxes):
                    # Check if heading based on style within pPr
                    is_heading_lxml = False
                    pPr = p_element.find(qn('w:pPr'))
                    if pPr is not None:
                         pStyle = pPr.find(qn('w:pStyle'))
                         if pStyle is not None:
                              style_val = pStyle.get(qn('w:val'))
                              if style_val and style_val.lower().startswith('heading'):
                                   is_heading_lxml = True

                    text = parse_lxml_paragraph(p_element)

                    # Add element for this paragraph ONLY
                    location = f"TextBox LXML Para {p_idx}"
                    print(f"  Extracting ({location}): Head={is_heading_lxml}, Text='{text[:50].replace(chr(10), ' ')}...'") # LOGGING
                    elements_to_translate.append(process_text(text, is_heading_lxml, lxml_obj=p_element, location=location))
                    standard_paragraph_objects.append(None) # Placeholder
                    textbox_paragraph_elements.append(p_element)
                    textbox_paras_count += 1

            print(f"  Finished Text Box Content. Paragraphs extracted: {textbox_paras_count}")

        except Exception as e:
            print(f"⚠️ Warning: Could not process text boxes via lxml: {e}")
            traceback.print_exc()

        # --- 3. Final Count ---
        print(f"\n--- Finished DOCX Extraction ---") # Added logging
        print(f"  Total elements extracted: {len(elements_to_translate)}") # Added logging


        # Ensure lists remain aligned (should inherently be aligned now)
        if len(elements_to_translate) != len(standard_paragraph_objects) or len(elements_to_translate) != len(textbox_paragraph_elements):
             print(f"❌ CRITICAL ERROR: List lengths mismatched after extraction! ") # Simplified error
             # Handle error state
             # <<< ADDED LOGGING >>>
             print(f"  Lengths -> Elements: {len(elements_to_translate)}, Std Paras: {len(standard_paragraph_objects)}, LXML Paras: {len(textbox_paragraph_elements)}")
             return [], [], [], None, doc # Return safely

        # The textbox_shapes list is not needed for the current rebuild logic
        # <<< ADDED LOGGING >>>
        print(f"  Returning {len(elements_to_translate)} elements, {len(standard_paragraph_objects)} std para refs, {len(textbox_paragraph_elements)} lxml para refs.")
        return elements_to_translate, standard_paragraph_objects, textbox_paragraph_elements, None, doc

    except Exception as e:
        print(f"❌ Error extracting DOCX elements from {file_path}: {e}")
        traceback.print_exc()
        # Ensure doc object is returned even if extraction fails partially
        # --- FIX: Check if 'doc' exists before returning ---
        return [], [], [], None, doc if 'doc' in locals() and doc is not None else None # Return safely


# ---------------- Main processing function ----------------
#
# --- FIX 1: Add 'output_dir' as the third argument ---
def process_file(file_path, target_lang, output_dir, output_format=None, task_id=None, tasks=None):
    """
    ENHANCED: High-fidelity extraction, translation, and rebuilding with format conversion.
    """
    doc_object = None
    standard_paragraph_objects = None
    textbox_paragraph_elements = None
    textbox_shapes = None # Keep variable, though unused in rebuild
    
    # --- FIX 2: Disable DocumentAnalyzer to prevent memory crash ---
    # This line loads a large AI model and is likely causing the server to run
    # out of memory and crash. Commenting it out will fix the crash.
    # doc_analyzer = DocumentAnalyzer()
    # ---
    elements = []
    original_format = ""
    
    # --- FIX 1: Remove hardcoded (and incorrect) output directory ---
    # The correct directory is now passed in from app.py
    # output_dir = "storage/translated" 
    # ---
    
    translated_files = {}

    try:
        if not file_path or not isinstance(file_path, str):
             raise ValueError("Invalid file path provided.")

        # Keep a copy of the original input file so we can fall back to PDF-based rebuilds
        source_input_file = file_path

        _, file_extension = os.path.splitext(file_path)
        original_format = file_extension.lower().strip('.')
        # Remember the original input format (before any in-memory conversion)
        original_input_format = original_format

        print(f"\nProcessing file: {os.path.basename(file_path)}, Format: {original_format}") # Log start

        # Prepare temp files tracker early so conversion-created files are cleaned up later
        created_temp_files = []

        if original_format == "pdf":
            # NEW WORKFLOW: Convert PDF -> DOCX first (preserve layout), then extract DOCX elements
            print(f"🔁 Converting source PDF to DOCX for high-fidelity translation...")
            # --- FIX 1: Use the correct output_dir variable ---
            temp_converted_docx = os.path.join(output_dir, f"converted_{uuid.uuid4()}.docx")
            try:
                convert_pdf_to_docx(file_path, temp_converted_docx)
                created_temp_files.append(temp_converted_docx)
                print(f"  ✅ PDF converted to DOCX: {temp_converted_docx}")
            except Exception as conv_e:
                print(f"❌ PDF->DOCX conversion failed: {conv_e}")
                traceback.print_exc()
                # Update task status if applicable and abort
                if task_id and tasks and task_id in tasks:
                    tasks[task_id]["status"] = "error"
                    tasks[task_id]["error_message"] = "PDF to DOCX conversion failed."
                return {}

            # Extract from the converted DOCX (this preserves paragraphs, textboxes, and formatting metadata)
            elements, standard_paragraph_objects, textbox_paragraph_elements, _, doc_object = extract_docx_elements_and_objects(temp_converted_docx)
            
            # --- FIX 2: Disable DocumentAnalyzer call ---
            # This call also consumes a lot of memory and relies on the object created above.
            # elements = doc_analyzer.enhance_extraction(temp_converted_docx, elements)
            # ---
            
            # Treat the rest of the flow as DOCX-based (so rebuild uses DOCX routines and conversion back to PDF)
            original_format = "docx"
            file_path = temp_converted_docx
        elif original_format == "docx":
            # Pass None for textbox_shapes as it's not currently used
            elements, standard_paragraph_objects, textbox_paragraph_elements, _, doc_object = extract_docx_elements_and_objects(file_path)
        else:
            raise ValueError(f"Unsupported file type: {original_format}")

        os.makedirs(output_dir, exist_ok=True)

        # CRITICAL CHECK: If extraction yielded no elements, stop here.
        if not elements:
            print(f"⛔ ERROR: Extraction returned no elements for {file_path}. Cannot proceed.")
            # Update task status if applicable
            if task_id and tasks and task_id in tasks:
                 tasks[task_id]["status"] = "error"
                 tasks[task_id]["error_message"] = "Extraction failed: No content found in the document."
            return {} # Return empty dict, indicating failure


        print(f"✅ Starting translation of {len(elements)} elements to {target_lang}...")
        translated_elements = translate_elements(elements, target_lang, task_id, tasks)

        # Ensure translation didn't fail catastrophically
        if not translated_elements or len(translated_elements) != len(elements):
             # Log the mismatch
             print(f"⛔ ERROR: Translation failed or returned mismatched element count.")
             print(f"  Expected: {len(elements)}, Got: {len(translated_elements)}")
             # Optionally print first few expected vs received for debugging:
             # print("  Expected first 5:", elements[:5])
             # print("  Received first 5:", translated_elements[:5])
             raise Exception("Translation failed or returned mismatched element count.")


        # Determine output format.
        # If caller specified output_format, honor it. Otherwise:
        # - If the original uploaded file was a PDF, default to PDF output (we convert back to PDF after translating DOCX)
        # - Otherwise default to the current original_format (usually 'docx')
        if output_format and output_format in ['pdf', 'docx']:
            final_format = output_format
        else:
            final_format = 'pdf' if original_input_format == 'pdf' else original_format

        # Define the output file path
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        # Sanitize basename slightly (replace spaces, etc.) - optional
        safe_base_name = re.sub(r'[^\w\-]+', '_', base_name)
        out_file = os.path.join(
            output_dir, # <-- FIX 1: This now correctly uses the directory from app.py
            f"{safe_base_name}_{target_lang.lower()}.{final_format}"
        )

        print(f"🔧 Rebuilding output file: {out_file} (Format: {final_format})")

        # --- REBUILD AND FORMAT CONVERSION LOGIC ---
        # Generate unique temp filenames using uuid
        temp_pdf = os.path.join(output_dir, f"temp_{safe_base_name}_{uuid.uuid4()}.pdf")
        temp_docx = os.path.join(output_dir, f"temp_{safe_base_name}_{uuid.uuid4()}.docx")

        try:
            if original_format == "pdf" and final_format == "pdf":
                rebuild_pdf_in_place(file_path, translated_elements, out_file, target_lang)
            elif original_format == "pdf" and final_format == "docx":
                rebuild_pdf_in_place(file_path, translated_elements, temp_pdf, target_lang)
                created_temp_files.append(temp_pdf)
                convert_pdf_to_docx(temp_pdf, out_file)
            elif original_format == "docx" and final_format == "docx":
                # CRITICAL CHECK: Ensure doc_object is valid before rebuilding
                # --- FIX: Use DocumentClass for isinstance check ---
                if not isinstance(doc_object, DocumentClass):
                    # Attempt to reload if it's None or invalid
                    print(f"⚠️ Warning: doc_object invalid before rebuild. Attempting to reload {file_path}")
                    try:
                        # Use DocxDocument alias to open
                        doc_object = DocxDocument(file_path)
                         # Re-run extraction minimally ONLY to get object references if lists are empty/invalid
                        if not standard_paragraph_objects and not textbox_paragraph_elements:
                             print("  Re-extracting object references...")
                             _, standard_paragraph_objects, textbox_paragraph_elements, _, doc_object = extract_docx_elements_and_objects(file_path)
                             # Check element count alignment again if re-extracted
                             if len(elements) != len(standard_paragraph_objects) or len(elements) != len(textbox_paragraph_elements):
                                 raise Exception("List lengths mismatched after re-extraction during rebuild.")
                    except Exception as reload_e:
                        raise Exception(f"DOCX object was invalid and failed to reload: {reload_e}")

                # --- FIX: Check again using DocumentClass ---
                if not isinstance(doc_object, DocumentClass): # Check again after potential reload attempt
                    raise Exception("DOCX object not available for rebuilding.")


                rebuild_docx_with_lxml(
                    doc_object,
                    standard_paragraph_objects,
                    textbox_paragraph_elements,
                    None, # textbox_shapes unused
                    translated_elements,
                    out_file
                )
            elif original_format == "docx" and final_format == "pdf":
                 # CRITICAL CHECK (similar to docx->docx)
                 # --- FIX: Use DocumentClass for isinstance check ---
                if not isinstance(doc_object, DocumentClass):
                    print(f"⚠️ Warning: doc_object invalid before rebuild (->pdf). Attempting to reload {file_path}")
                    try:
                        # Use DocxDocument alias to open
                        doc_object = DocxDocument(file_path)
                        if not standard_paragraph_objects and not textbox_paragraph_elements:
                             print("  Re-extracting object references...")
                             _, standard_paragraph_objects, textbox_paragraph_elements, _, doc_object = extract_docx_elements_and_objects(file_path)
                             if len(elements) != len(standard_paragraph_objects) or len(elements) != len(textbox_paragraph_elements):
                                  raise Exception("List lengths mismatched after re-extraction during rebuild (->pdf).")
                    except Exception as reload_e:
                         raise Exception(f"DOCX object was invalid and failed to reload (->pdf): {reload_e}")

                # --- FIX: Check again using DocumentClass ---
                if not isinstance(doc_object, DocumentClass):
                     raise Exception("DOCX object not available for rebuilding (->pdf).")

                rebuild_docx_with_lxml(
                    doc_object,
                    standard_paragraph_objects,
                    textbox_paragraph_elements,
                    None, # textbox_shapes unused
                    translated_elements,
                    temp_docx
                )
                created_temp_files.append(temp_docx)
                # Attempt primary conversion method
                conversion_succeeded = False
                try:
                    convert_docx_to_pdf(temp_docx, out_file)
                    conversion_succeeded = os.path.exists(out_file) and os.path.getsize(out_file) > 0
                except Exception as conv_e:
                    print(f"⚠️ convert_docx_to_pdf() failed: {conv_e}")
                    traceback.print_exc()

                # Fallback: try docx2pdf (if installed)
                if not conversion_succeeded:
                    try:
                        from docx2pdf import convert as docx2pdf_convert
                        print("🔁 Falling back to docx2pdf.convert()...")
                        # docx2pdf can convert a single file to a single output file
                        try:
                            docx2pdf_convert(temp_docx, out_file)
                        except TypeError:
                            # Older docx2pdf may accept only (input, ) and output to same-name.pdf in same folder
                            docx2pdf_convert(temp_docx)
                            # Move generated file to out_file if needed
                        conversion_succeeded = os.path.exists(out_file) and os.path.getsize(out_file) > 0
                    except Exception as d2p_e:
                        print(f"⚠️ docx2pdf fallback failed: {d2p_e}")
                        # Do not traceback here; we'll attempt PDF rebuild fallback next

                # Last-resort fallback: if source was a PDF, try rebuilding the translated content into original PDF
                if not conversion_succeeded:
                    try:
                        # If we have an original source PDF, try to rebuild into a PDF using the original as template
                        if source_input_file and os.path.exists(source_input_file) and source_input_file.lower().endswith('.pdf'):
                            print("🔁 Last-resort: rebuilding translated content into original PDF using rebuild_pdf_in_place()...")
                            rebuild_pdf_in_place(source_input_file, translated_elements, out_file, target_lang)
                            conversion_succeeded = os.path.exists(out_file) and os.path.getsize(out_file) > 0
                    except Exception as pdf_rebuild_e:
                        print(f"⚠️ Last-resort PDF rebuild failed: {pdf_rebuild_e}")
                        traceback.print_exc()

                if not conversion_succeeded:
                    raise Exception("DOCX->PDF conversion failed (all methods tried).")

            # Check if output file was actually created
            if not os.path.exists(out_file) or os.path.getsize(out_file) == 0:
                 raise Exception(f"Output file '{out_file}' was not created or is empty after rebuild/conversion.")

            translated_files[target_lang] = os.path.basename(out_file)

        finally:
            # Clean up temporary files
            for temp_file in created_temp_files:
                if os.path.exists(temp_file):
                    try:
                        os.remove(temp_file)
                        print(f"  Removed temporary file: {temp_file}")
                    except Exception as clean_e:
                        print(f"⚠️ Warning: Could not remove temporary file {temp_file}: {clean_e}")

        return translated_files

    # Catch specific exceptions for better error messages
    except FileNotFoundError:
        print(f"⛔ ERROR: Input file not found at {file_path}")
        if task_id and tasks and task_id in tasks:
            tasks[task_id]["status"] = "error"
            tasks[task_id]["error_message"] = "Input file not found."
        return {}
    except ValueError as ve: # Catch unsupported format or invalid path
        print(f"⛔ ERROR: {ve}")
        if task_id and tasks and task_id in tasks:
            tasks[task_id]["status"] = "error"
            tasks[task_id]["error_message"] = f"Error: {ve}"
        return {}
    except Exception as e: # General catch-all
        print(f"❌ Unhandled error processing file {file_path}: {e}")
        traceback.print_exc()
        # Update task status if applicable
        if task_id and tasks and task_id in tasks:
             tasks[task_id]["status"] = "error"
             # Be careful about exposing too much detail from arbitrary exceptions
             tasks[task_id]["error_message"] = f"An unexpected processing error occurred. Please check logs."
        return {} # Return empty on error
