"""
Document Layout Analysis Module
Enhances document processing by analyzing layout and structure.
"""

import os
from transformers import LayoutLMModel, LayoutLMTokenizer
import numpy as np
from PIL import Image
import torch
from pdf2docx import Converter
from docx import Document
import fitz  # PyMuPDF

class DocumentAnalyzer:
    def __init__(self):
        self.model_path = os.path.join('models', 'layout', 'layoutlm-base-uncased')
        self.tokenizer = None
        self.model = None
        self._load_models()

    def _load_models(self):
        """Load LayoutLM model and tokenizer."""
        try:
            if not os.path.exists(self.model_path):
                print("Downloading LayoutLM model (first time only)...")
                self.tokenizer = LayoutLMTokenizer.from_pretrained('microsoft/layoutlm-base-uncased')
                self.model = LayoutLMModel.from_pretrained('microsoft/layoutlm-base-uncased')
                
                # Save models locally
                os.makedirs(self.model_path, exist_ok=True)
                self.tokenizer.save_pretrained(self.model_path)
                self.model.save_pretrained(self.model_path)
            else:
                self.tokenizer = LayoutLMTokenizer.from_pretrained(self.model_path)
                self.model = LayoutLMModel.from_pretrained(self.model_path)
            
            print("✅ LayoutLM model loaded successfully")
        except Exception as e:
            print(f"⚠️ Warning: Could not load LayoutLM model: {e}")
            print("Document analysis will proceed with basic layout detection")
    
    def analyze_document_layout(self, file_path):
        """
        Analyze document layout to improve extraction accuracy.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            dict: Layout analysis information
        """
        if file_path.lower().endswith('.pdf'):
            return self._analyze_pdf_layout(file_path)
        elif file_path.lower().endswith('.docx'):
            return self._analyze_docx_layout(file_path)
        else:
            raise ValueError(f"Unsupported file format: {os.path.splitext(file_path)[1]}")
    
    def _analyze_pdf_layout(self, pdf_path):
        """
        Analyze PDF document layout.
        """
        layout_info = {
            'pages': [],
            'structure': {
                'headers': [],
                'body_blocks': [],
                'tables': [],
                'images': [],
                'footnotes': []
            }
        }
        
        try:
            doc = fitz.open(pdf_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_dict = page.get_text("dict")
                
                # Analyze page structure
                page_info = {
                    'number': page_num + 1,
                    'size': page.rect.br,  # Bottom-right point gives page dimensions
                    'blocks': []
                }
                
                # Process text blocks
                for block in page_dict.get("blocks", []):
                    if block["type"] == 0:  # Text block
                        block_info = {
                            'type': 'text',
                            'bbox': block["bbox"],
                            'text': ' '.join(word["text"] for line in block["lines"] 
                                           for span in line["spans"] 
                                           for word in span["text"].split()),
                            'font_size': block["lines"][0]["spans"][0]["size"] 
                                       if block["lines"] and block["lines"][0]["spans"] else None
                        }
                        page_info['blocks'].append(block_info)
                        
                        # Classify block based on position and properties
                        if block_info['bbox'][1] < 100:  # Near top of page
                            layout_info['structure']['headers'].append(block_info)
                        elif block_info['bbox'][3] > page.rect.height - 100:  # Near bottom
                            layout_info['structure']['footnotes'].append(block_info)
                        else:
                            layout_info['structure']['body_blocks'].append(block_info)
                    
                    elif block["type"] == 1:  # Image block
                        layout_info['structure']['images'].append({
                            'type': 'image',
                            'bbox': block["bbox"]
                        })
                
                layout_info['pages'].append(page_info)
            
            return layout_info
        
        except Exception as e:
            print(f"⚠️ Warning: Layout analysis error: {e}")
            return layout_info
    
    def _analyze_docx_layout(self, docx_path):
        """
        Analyze DOCX document layout.
        """
        layout_info = {
            'sections': [],
            'structure': {
                'headers': [],
                'body_paragraphs': [],
                'tables': [],
                'text_boxes': []
            }
        }
        
        try:
            doc = Document(docx_path)
            
            # Analyze sections
            for section in doc.sections:
                section_info = {
                    'orientation': 'portrait' if section.orientation == 0 else 'landscape',
                    'width': section.page_width.pt,
                    'height': section.page_height.pt
                }
                layout_info['sections'].append(section_info)
            
            # Analyze paragraphs
            for para in doc.paragraphs:
                para_info = {
                    'style': para.style.name if para.style else 'Normal',
                    'text': para.text,
                    'runs': [{'text': run.text, 'bold': run.bold, 'italic': run.italic}
                            for run in para.runs]
                }
                
                # Classify based on style
                if para.style and 'heading' in para.style.name.lower():
                    layout_info['structure']['headers'].append(para_info)
                else:
                    layout_info['structure']['body_paragraphs'].append(para_info)
            
            # Analyze tables
            for table in doc.tables:
                table_info = {
                    'rows': len(table.rows),
                    'cols': len(table.columns),
                    'cells': []
                }
                for row in table.rows:
                    for cell in row.cells:
                        table_info['cells'].append({
                            'text': cell.text,
                            'paragraphs': len(cell.paragraphs)
                        })
                layout_info['structure']['tables'].append(table_info)
            
            return layout_info
        
        except Exception as e:
            print(f"⚠️ Warning: DOCX layout analysis error: {e}")
            return layout_info

    def enhance_extraction(self, file_path, extracted_elements):
        """
        Enhance extracted elements using layout analysis.
        
        Args:
            file_path: Path to the document file
            extracted_elements: List of previously extracted elements
            
        Returns:
            list: Enhanced extracted elements
        """
        try:
            layout_info = self.analyze_document_layout(file_path)
            enhanced_elements = []
            
            for element in extracted_elements:
                # Enhance element with layout context
                enhanced = element.copy()
                
                # Add structural context
                if element.get('is_heading', False):
                    enhanced['structural_role'] = 'header'
                elif element.get('is_equation', False):
                    enhanced['structural_role'] = 'equation'
                else:
                    # Determine role based on layout analysis
                    enhanced['structural_role'] = self._determine_structural_role(
                        element, layout_info
                    )
                
                enhanced_elements.append(enhanced)
            
            return enhanced_elements
        
        except Exception as e:
            print(f"⚠️ Warning: Enhancement failed: {e}")
            return extracted_elements
    
    def _determine_structural_role(self, element, layout_info):
        """
        Determine the structural role of an element based on layout analysis.
        """
        # Default role
        role = 'body'
        
        try:
            text = element.get('text', '')
            
            # Check headers
            for header in layout_info['structure']['headers']:
                if isinstance(header, dict) and header.get('text') == text:
                    return 'header'
            
            # Check footnotes (PDF only)
            if 'footnotes' in layout_info['structure']:
                for footnote in layout_info['structure']['footnotes']:
                    if isinstance(footnote, dict) and footnote.get('text') == text:
                        return 'footnote'
            
            # More specific roles can be determined based on layout_info
            
        except Exception as e:
            print(f"⚠️ Warning: Role determination failed: {e}")
        
        return role