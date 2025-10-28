import shutil
import re
import fitz  # Import PyMuPDF (fitz)
import os    # Import os for path joining
from docx import Document
# FIX: Import the specific class type for isinstance checks
from docx.document import Document as DocumentClass
from docx.shared import Pt, RGBColor # Import Pt and RGBColor for formatting
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # Import alignment constants
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
# pdf2docx is imported conditionally later
# comtypes.client is imported conditionally later
import subprocess
import platform
import traceback # Import traceback for detailed error logging
# NEW: Import deepcopy for safer template cloning
from copy import deepcopy


# Compile the regex pattern for bold
bold_pattern = re.compile(r'\*\*(.*?)\*\*')

# --- Map language to expected TTF FILENAME ---
# Assumes these files exist in a 'fonts' directory relative to this script's location
LANGUAGE_TTF_MAP = {
    "Hindi": "NotoSansDevanagari-Regular.ttf",
    "Marathi": "NotoSansDevanagari-Regular.ttf",
    "Bengali": "NotoSansBengali-Regular.ttf",
    "Japanese": "NotoSansJP-Regular.otf",
    "Chinese": "NotoSansSC-Regular.otf",
    "Irish": "NotoSansIrish-Regular.ttf",
    "Tamil": "NotoSansTamil-Regular.ttf",
    "Italian": "NotoSansItalic-Regular.ttf",
    "Spanish": "NotoSans-Regular.ttf",
    "French": "NotoSans-Regular.ttf",
    "German": "NotoSans-Regular.ttf",
    "Russian": "NotoSans-Regular.ttf",
    "Arabic": "NotoSansArabic-Regular.ttf",
    "Telugu": "NotoSansTelugu-Regular.ttf",
    "Canada": "NotoSans-Regular.ttf",
}
DEFAULT_FONT_REF = "helv" # PyMuPDF built-in reference name for Helvetica


# ---------------- PDF ----------------
def rebuild_pdf_in_place(original_file_path, translated_elements, output_path, target_lang):
    """
    ENHANCED: Rebuilds PDF preserving ALL images, applying font styles (bold/italic),
    and maintaining exact layout. Uses fontfile parameter directly.
    """
    doc = None
    # REMOVED font buffers, will read paths
    regular_font_buffer = None
    bold_font_buffer = None

    try:
        # 1. Open the original PDF
        doc = fitz.open(original_file_path)
        print(f"--- Starting PDF Rebuild for: {os.path.basename(output_path)} ---") # Logging

        # --- Prepare font file paths (regular, bold, italic) ---
        target_ttf = LANGUAGE_TTF_MAP.get(target_lang)
        font_path = None
        bold_font_path = None
        italic_font_path = None # REMOVED: Italic path removed as requested
        regular_font_found = False # Flag to track if the main font file exists

        print(f"\n🔍 Font Setup for Target Language: '{target_lang}'")
        print(f"  Mapped font filename: {target_ttf if target_ttf else 'None (using default)'}")

        if target_ttf:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            project_root = os.path.dirname(script_dir) # Go up one level from src
            fonts_dir = os.path.join(project_root, "fonts") # Path to fonts directory
            
            print(f"  Script directory: {script_dir}")
            print(f"  Project root: {project_root}")
            print(f"  Fonts directory: {fonts_dir}")
            print(f"  Fonts directory exists: {os.path.exists(fonts_dir)}")

            font_path = os.path.join(fonts_dir, target_ttf)
            base_name, ext = os.path.splitext(target_ttf)
            # Construct potential bold/italic filenames
            bold_ttf = f"{base_name.replace('Regular', 'Bold')}{ext}"
            # italic_ttf = f"{base_name.replace('Regular', 'Italic')}{ext}" # REMOVED
            bold_font_path = os.path.join(fonts_dir, bold_ttf)
            # italic_font_path = os.path.join(fonts_dir, italic_ttf) # REMOVED

            # Verify font files exist and update logs
            if os.path.exists(font_path):
                print(f"  ✅ Found regular font file: '{font_path}'")
                print(f"  Using regular font: '{os.path.basename(font_path)}'")
                regular_font_found = True # Set flag
                # Check for bold variant
                if os.path.exists(bold_font_path):
                    print(f"  ✅ Found bold font file: '{bold_font_path}'")
                    print(f"  Using bold font: '{os.path.basename(bold_font_path)}'")
                else:
                    bold_font_path = None # Mark as not found
                    print(f"  ⚠️ Bold font variant not found: {bold_ttf}. Will use regular.")
                # REMOVED: Italic check removed
            else:
                print(f"❌ Regular font file NOT FOUND at: '{font_path}'")
                # Try to find an appropriate system-installed font as a fallback (Windows, macOS, Linux)
                def _search_system_fonts(candidates=None):
                    # candidates: list of lowercase substrings to look for in font filenames
                    cand = candidates or []
                    sys_paths = []
                    system = platform.system()
                    if system == "Windows":
                        sys_paths = [r"C:\\Windows\\Fonts"]
                    elif system == "Darwin":
                        sys_paths = ["/Library/Fonts", "/System/Library/Fonts"]
                    else:
                        sys_paths = ["/usr/share/fonts", "/usr/local/share/fonts", os.path.expanduser("~/.local/share/fonts")]

                    for p in sys_paths:
                        try:
                            if not os.path.exists(p):
                                continue
                            for root, dirs, files in os.walk(p):
                                for f in files:
                                    fn = f.lower()
                                    if fn.endswith((".ttf", ".otf")):
                                        # If any candidate substring appears in filename, return it
                                        for sub in cand:
                                            if sub in fn:
                                                return os.path.join(root, f)
                        except Exception:
                            continue
                    return None

                # Common substrings for Devanagari / Indic fonts and Noto family
                search_subs = []
                # include the requested ttf filename parts
                if target_ttf:
                    search_subs.append(os.path.splitext(target_ttf)[0].lower())
                # common Windows/Unicode fonts
                search_subs += ["mangal", "nirmala", "noto", "devanagari", "sansdevanagari"]

                found_sys_font = _search_system_fonts(search_subs)
                if found_sys_font:
                    font_path = found_sys_font
                    regular_font_found = True
                    print(f"  ✅ Found system fallback font: {font_path}")
                else:
                    print(f"⚠️ Falling back to default font. Hindi/non-English text will NOT display correctly!")
                    font_path = None; bold_font_path = None

        # --- Process each page ---
        for page_num in range(doc.page_count):
            page = doc[page_num]
            print(f"  Rebuilding PDF Page {page_num + 1}/{doc.page_count}") # Logging
            page_elements = sorted(
                [el for el in translated_elements if el.get("page_num") == page_num],
                key=lambda el: (el.get("bbox", (0, 0, 0, 0))[1], el.get("bbox", (0, 0, 0, 0))[0]) # Sort top-to-bottom, left-to-right
            )
            if not page_elements: continue

            # --- PRE-EMBED FONTS FOR THIS PAGE (CRITICAL FIX) ---
            regular_fontname = None
            bold_fontname = None
            
            print(f"\n  🔧 Font Embedding for Page {page_num + 1}:")
            print(f"    regular_font_found: {regular_font_found}")
            print(f"    font_path: {font_path}")
            print(f"    bold_font_path: {bold_font_path}")
            
            # Pre-embed regular font
            if regular_font_found and font_path:
                try:
                    print(f"    🔴 Attempting to embed: {font_path}")
                    xref = page.insert_font(fontfile=font_path, encoding=0)  # Returns integer xref
                    regular_fontname = f"/{xref}"  # Convert to string format "/48", "/49", etc.
                    print(f"  ✅ Pre-embedded regular font: {os.path.basename(font_path)} -> {regular_fontname}")
                except Exception as font_e:
                    print(f"  ❌ Failed to pre-embed regular font: {font_e}")
                    traceback.print_exc()
                    regular_fontname = None
            else:
                print(f"  ⚠️ Skipping regular font embedding (not found or flag false)")
            
            # Pre-embed bold font
            if bold_font_path:
                try:
                    print(f"    🔴 Attempting to embed bold: {bold_font_path}")
                    xref_bold = page.insert_font(fontfile=bold_font_path, encoding=0)
                    bold_fontname = f"/{xref_bold}"
                    print(f"  ✅ Pre-embedded bold font: {os.path.basename(bold_font_path)} -> {bold_fontname}")
                except Exception as bold_e:
                    print(f"  ❌ Failed to pre-embed bold font: {bold_e}")
                    traceback.print_exc()
                    bold_fontname = None
            else:
                print(f"  🚨 No bold font path available")

            # --- Redact original text areas ---
            image_rects = []
            try: # Defensive coding for image extraction
                for img_info in page.get_images(full=True):
                    try:
                        bboxes = page.get_image_bbox(img_info)
                        if bboxes:
                            img_rect = fitz.Rect(bboxes if isinstance(bboxes, tuple) else bboxes[0])
                            image_rects.append(img_rect)
                    except Exception as img_e: pass # Ignore bbox errors
            except Exception as get_img_e:
                 print(f"  ⚠️ Error getting images for redaction check: {get_img_e}")

            redactions_added = 0 # Logging
            for el in page_elements:
                 if el.get("text", "[EMPTY]") != "[EMPTY]" and "bbox" in el:
                    try:
                        text_rect = fitz.Rect(el["bbox"])
                        if text_rect.is_empty or text_rect.width <= 0 or text_rect.height <= 0: continue # Skip invalid rects
                        overlaps_image = False
                        if image_rects:
                            for img_rect in image_rects:
                                intersection = text_rect & img_rect
                                # Check if intersection area is significant compared to text area
                                if not intersection.is_empty and (intersection.width * intersection.height) > (text_rect.width * text_rect.height * 0.5):
                                    overlaps_image = True; break
                        if not overlaps_image:
                            page.add_redact_annot(text_rect, fill=(1, 1, 1)); redactions_added += 1
                    except Exception as redact_e: print(f"  Error adding redaction for bbox {el.get('bbox')}: {redact_e}")

            try:
                if redactions_added > 0:
                     page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE) # type: ignore
            except Exception as apply_redact_e: print(f"  Error applying redactions: {apply_redact_e}")


            # --- Insert translated text ---
            insertions_done = 0 # Logging
            for el in page_elements:
                text = el.get("text", "[EMPTY]")
                if text == "[EMPTY]" or "bbox" not in el: continue
                try:
                    bbox = fitz.Rect(el["bbox"])
                    if bbox.is_empty or bbox.width <= 0 or bbox.height <= 0:
                        print(f"  ⚠️ Skipping element - invalid or empty bbox dimensions: {bbox}")
                        continue

                    # Determine which FONT REFERENCE to use (from pre-embedded fonts)
                    is_bold_in_el = el.get("is_bold", False)

                    # --- Font Selection Logic (using pre-embedded font references) ---
                    selected_fontname = None  # Will be like "/48", "/49" from insert_font()
                    
                    if is_bold_in_el and bold_fontname:
                        selected_fontname = bold_fontname  # Use pre-embedded bold font
                    elif regular_fontname:
                        selected_fontname = regular_fontname  # Use pre-embedded regular font
                    else:
                        # Fallback to built-in font if no custom fonts were embedded
                        selected_fontname = DEFAULT_FONT_REF  # "helv" (Helvetica)
                        if page_num == 0 and insertions_done < 1:
                            print(f"  ⚠️ No custom fonts embedded, using default '{selected_fontname}' for page {page_num + 1}")


                    text_to_render = text.replace("**", "") # Remove bold markers if any

                    # Get alignment
                    align = fitz.TEXT_ALIGN_LEFT # Default to left
                    align_str = el.get("alignment", "LEFT").upper()
                    if align_str == "CENTER": align = fitz.TEXT_ALIGN_CENTER
                    elif align_str == "RIGHT": align = fitz.TEXT_ALIGN_RIGHT
                    elif align_str == "JUSTIFY": align = fitz.TEXT_ALIGN_JUSTIFY

                    font_size = max(el.get("font_size", 9.0), 4.0) # Ensure min size 4

                    # Try inserting text with automatic font size reduction if needed
                    result = -1 # PyMuPDF returns negative value if text doesn't fit
                    current_font_size = font_size
                    min_font_size = 4.0 # Minimum readable font size

                    while result < 0 and current_font_size >= min_font_size:
                        try:
                            # ✅ FIXED: Use pre-embedded font reference (string like "/48")
                            result = page.insert_textbox(
                                bbox,
                                text_to_render,
                                fontsize=current_font_size,
                                fontname=selected_fontname,  # ✅ Use pre-embedded font reference
                                align=align,
                                color=(0, 0, 0)  # Black text
                            )

                        except Exception as textbox_e:
                            print(f"  ❌ Error during insert_textbox: {textbox_e}")
                            print(f"     Text: '{text_to_render[:50]}...', Font: {selected_fontname}, Size: {current_font_size}")
                            traceback.print_exc()
                            result = -99 # Indicate failure due to exception
                            break # Stop trying for this element

                        if result < 0 and current_font_size >= min_font_size + 1: # Only reduce if not already at min
                            current_font_size -= 1.0 # Decrement by 1 point
                        elif result >= 0:
                            # Success!
                            if current_font_size < font_size:
                                print(f"  📝 Reduced font size from {font_size:.1f} to {current_font_size:.1f} for better fit for '{text_to_render[:20]}...'")
                            insertions_done += 1
                            break # Exit the while loop on success
                        else: # result < 0 and current_font_size < min_font_size + 1
                             break # Stop trying if already at minimum or below


                    # Log failure if result is still negative after trying
                    if result < 0:
                        print(f"  ⚠️ Text insertion failed even at minimum font size {min_font_size:.1f}: '{text_to_render[:50]}...'")
                        print(f"     Original size: {font_size:.1f}, BBox: {bbox}, Font: {selected_fontname}")

                except Exception as insert_e:
                    print(f"  ❌ Error processing element to insert text '{text[:50]}...': {insert_e}")
                    traceback.print_exc()

            print(f"    ✅ Inserted {insertions_done} text elements on page {page_num + 1}")

        # Save the modified document
        doc.save(output_path, garbage=3, deflate=True, clean=True)
        print(f"✅ PDF rebuilt successfully: {output_path}") # Logging

    except Exception as e:
        print(f"❌ PDF rebuild failed: {e}")
        traceback.print_exc()
    finally:
        if doc:
            try: doc.close()
            except: pass


# ---------------- DOCX (DEPRECATED - Use LXML version) ----------------
# Keep this function definition for compatibility, but it won't be used by default
# Or add a stronger warning/raise an error if called.
def rebuild_docx_in_place(original_file_path, translated_elements, output_path):
    """ DEPRECATED - Use rebuild_docx_with_lxml for better results """
    print("⚠️ WARNING: Calling DEPRECATED rebuild_docx_in_place function. Use rebuild_docx_with_lxml instead.")
    doc = None
    try:
        shutil.copy(original_file_path, output_path)
        doc = Document(output_path)
        para_index = 0
        if not isinstance(translated_elements, (list, tuple)):
             raise TypeError("translated_elements must be a list or tuple")
        for i, el in enumerate(translated_elements):
             if el.get("type") == "paragraph":
                 if para_index < len(doc.paragraphs):
                     original_para = doc.paragraphs[para_index]
                     original_style = original_para.style
                     original_alignment = original_para.alignment
                     orig_fmt = original_para.paragraph_format # Renamed variable

                     # --- Snapshotting run format (Improved) ---
                     template_run_props = {'bold': None, 'italic': None, 'font_name': None, 'font_size': None, 'font_color': None}
                     if original_para.runs:
                          first_run = original_para.runs[0]
                          template_run_props['bold'] = getattr(first_run, 'bold', None) # Use getattr for safety
                          template_run_props['italic'] = getattr(first_run, 'italic', None)
                          if hasattr(first_run.font, 'name'): template_run_props['font_name'] = first_run.font.name
                          if hasattr(first_run.font, 'size'): template_run_props['font_size'] = first_run.font.size # Might be None or Pt
                          if hasattr(first_run.font, 'color') and hasattr(first_run.font.color, 'rgb'):
                               template_run_props['font_color'] = first_run.font.color.rgb # Might be None or RGBColor
                     # --- End Snapshot ---


                     # Clear content more safely
                     for run in list(original_para.runs):
                          try: original_para._element.remove(run._element)
                          except ValueError: pass # Ignore if already removed

                     text = el.get("text", "")
                     if text != "[EMPTY]":
                        parts = bold_pattern.split(text)

                        for j, part in enumerate(parts):
                            if part:
                                run = original_para.add_run(part)
                                is_bold_marker = (j % 2 == 1)

                                # Apply formatting from snapshot
                                # --- FIX for Linter Error ---
                                run.bold = is_bold_marker or (template_run_props.get('bold') is True) # Explicitly check True
                                run.italic = (template_run_props.get('italic') is True) # Explicitly check True, removed italic marker logic
                                
                                font_name = template_run_props.get('font_name')
                                if font_name is not None:
                                    run.font.name = font_name
                                
                                font_size = template_run_props.get('font_size')
                                if font_size is not None:
                                     try: run.font.size = font_size
                                     except: pass # Ignore size errors
                                
                                font_color = template_run_props.get('font_color')
                                if font_color is not None:
                                     try: run.font.color.rgb = font_color
                                     except: pass # Ignore color errors
                                # --- End Linter Fix ---

                     # Restore paragraph formatting
                     if original_style: original_para.style = original_style
                     if original_alignment is not None: original_para.alignment = original_alignment
                     curr_fmt = original_para.paragraph_format # Renamed variable
                     # Copy known format attributes defensively
                     for attr in ['left_indent', 'right_indent', 'first_line_indent', 'space_before', 'space_after', 'line_spacing', 'line_spacing_rule']:
                         if hasattr(orig_fmt, attr):
                            val = getattr(orig_fmt, attr) # Get value
                            if val is not None: # Check if it's not None
                                 try: setattr(curr_fmt, attr, val) # Set it
                                 except: pass # Ignore errors during set
                 para_index += 1
        doc.save(output_path)
        print(f"✅ DOCX saved (in-place - DEPRECATED METHOD) at {output_path}")
    except Exception as e:
        print(f"❌ Error during DEPRECATED DOCX in-place rebuild: {e}")
        traceback.print_exc()


# ---------------- NEW LXML HELPERS (for text boxes) ----------------

def _clear_lxml_paragraph(p_element):
    """
    Safely removes only text runs (<w:r> containing <w:t>) from an lxml paragraph element,
    PRESERVING paragraph properties (<w:pPr>).
    """
    runs_to_remove = []
    # Find all runs directly under the paragraph
    for r_element in p_element.xpath('w:r'): # More direct XPath
        # Remove if it contains text OR if it's just an empty run
        if r_element.xpath('w:t') or not list(r_element): # Check children count
             runs_to_remove.append(r_element)
    for r_element in runs_to_remove:
        try: p_element.remove(r_element)
        except ValueError: pass # Ignore if already gone

def _add_lxml_run(p_element, text, template_run_pr=None, is_bold=False):
    """ Adds a run (<w:r>) with text to an lxml paragraph, safely applying formatting. """
    r_element = OxmlElement('w:r')
    run_pr = OxmlElement('w:rPr') # Start with empty properties
    has_formatting = False
    if template_run_pr is not None:
        try:
            run_pr = deepcopy(template_run_pr) # Clone template properties
            has_formatting = True
            # Explicitly remove bold from template if the current part ISN'T bold
            if not is_bold:
                bold_tag = run_pr.find(qn('w:b'))
                if bold_tag is not None: run_pr.remove(bold_tag)
            # REMOVED: Italic logic
        except Exception as copy_e:
             run_pr = OxmlElement('w:rPr'); has_formatting = False # Fallback

    # Apply specific bold property if needed AND not already present from template
    if is_bold and run_pr.find(qn('w:b')) is None:
        b = OxmlElement('w:b'); run_pr.append(b); has_formatting = True
    
    # REMOVED: Italic logic

    # Append run properties ONLY if they contain actual formatting tags
    if has_formatting and list(run_pr): r_element.append(run_pr)

    # Add the text element
    t_element = OxmlElement('w:t'); t_element.text = text
    space_attr = qn('xml:space')
    # Preserve spaces if text contains leading/trailing spaces or multiple spaces
    if text and (text.strip() != text or "  " in text): # Add check for non-empty text
        t_element.set(space_attr, 'preserve')
    r_element.append(t_element)

    # Append run correctly (simplified): just append at the end
    p_element.append(r_element)


# ---------------- NEW LXML REBUILDER (for text boxes) ----------------

def rebuild_docx_with_lxml(doc, standard_paragraphs, textbox_paragraphs, textbox_shapes, translated_elements, output_path):
    """
    SIMPLIFIED: Rebuilds DOCX with straightforward index-based processing.
    Focus on reliability over perfection. Uses LXML for text boxes.
    """
    print(f"\n--- DOCX Rebuild (LXML Enhanced): {os.path.basename(output_path)} ---")

    if not translated_elements:
        print("⚠️ No translated elements provided. Saving empty document.")
        try:
            doc.save(output_path); print(f"✅ Empty document saved at {output_path}")
        except Exception as save_e: print(f"❌ Error saving empty doc: {save_e}")
        return

    print(f"  Processing {len(translated_elements)} elements...")
    processed_count = 0; skipped_count = 0

    # Assume lists are aligned from extraction
    total_items = len(translated_elements)
    # --- LINTER FIX: Check against DocumentClass ---
    if not (doc and isinstance(doc, DocumentClass) and
            standard_paragraphs is not None and len(standard_paragraphs) == total_items and
            textbox_paragraphs is not None and len(textbox_paragraphs) == total_items):
         print("❌ Error: Invalid input document object or list lengths mismatch (elements, std_paras, lxml_paras). Cannot rebuild.")
         return # Stop if inputs are invalid

    for i, translated_el in enumerate(translated_elements):
        translated_text = translated_el.get("text", "")
        if translated_text == "[EMPTY]": skipped_count += 1; continue

        std_para = standard_paragraphs[i] # Directly access by index
        lxml_para = textbox_paragraphs[i] # Directly access by index

        try:
            if std_para is not None:
                # --- Process standard paragraph (using python-docx) ---
                # Snapshot formatting
                original_style = std_para.style
                original_alignment = std_para.alignment
                orig_fmt = std_para.paragraph_format
                template_rpr = None # Run properties template
                if std_para.runs:
                    first_run = std_para.runs[0]
                    template_rpr = { # Use dict for easier handling
                        'bold': getattr(first_run, 'bold', None),
                        'italic': getattr(first_run, 'italic', None),
                        'font_name': getattr(first_run.font, 'name', None),
                        'font_size': getattr(first_run.font, 'size', None),
                        'font_color': getattr(first_run.font.color, 'rgb', None) if first_run.font.color else None
                    }

                # Clear paragraph content safely
                std_para.clear() # Should remove all runs

                # Restore paragraph level formatting
                if original_style:
                    try: std_para.style = original_style
                    except Exception as style_e: print(f"  Warn: Failed to apply style {original_style}: {style_e}")
                if original_alignment is not None:
                    try: std_para.alignment = original_alignment
                    except Exception as align_e: print(f"  Warn: Failed to apply alignment {original_alignment}: {align_e}")

                curr_fmt = std_para.paragraph_format
                for attr in ['space_before', 'space_after', 'left_indent', 'right_indent', 'first_line_indent', 'line_spacing', 'line_spacing_rule']:
                    try:
                        if hasattr(orig_fmt, attr):
                            val = getattr(orig_fmt, attr)
                            if val is not None: setattr(curr_fmt, attr, val)
                    except Exception as fmt_e: print(f"  Warn: Failed to apply format {attr}: {fmt_e}")


                # Add translated text runs
                parts = bold_pattern.split(translated_text)
                for j, part in enumerate(parts):
                    if part:
                        run = std_para.add_run(part)
                        is_bold_marker = (j % 2 == 1)
                        if template_rpr: # Apply template if available
                             # --- FIX for Linter Error ---
                             run.bold = is_bold_marker or (template_rpr.get('bold') is True) # Explicitly check True
                             run.italic = (template_rpr.get('italic') is True) # Explicitly check True, removed italic marker logic
                             
                             font_name = template_rpr.get('font_name')
                             if font_name is not None:
                                 run.font.name = font_name
                             
                             font_size = template_rpr.get('font_size')
                             if font_size is not None:
                                  try: run.font.size = font_size
                                  except: pass
                             
                             font_color = template_rpr.get('font_color')
                             if font_color is not None:
                                  try: run.font.color.rgb = font_color
                                  except: pass
                             # --- End Linter Fix ---
                        else: # Basic bold if no template
                             run.bold = is_bold_marker
                processed_count += 1

            elif lxml_para is not None:
                # --- Process text box paragraph (using lxml) ---
                template_rpr_lxml = None
                first_run_lxml = lxml_para.find(qn('w:r'))
                if first_run_lxml is not None: template_rpr_lxml = first_run_lxml.find(qn('w:rPr'))

                _clear_lxml_paragraph(lxml_para) # Clear existing runs

                parts = bold_pattern.split(translated_text)
                for j, part in enumerate(parts):
                    if part:
                        is_bold = (j % 2 == 1)
                        _add_lxml_run(lxml_para, part, template_run_pr=template_rpr_lxml, is_bold=is_bold)
                processed_count += 1
            else:
                # This case shouldn't happen if extraction logic ensures one is None and other is not
                print(f"  ⚠️ Element {i}: Neither standard para nor LXML para object found.")
                skipped_count += 1
        except Exception as para_e:
            print(f"  ❌ Error processing element {i}: {para_e}")
            traceback.print_exc() # Log full traceback for debugging
            skipped_count += 1

    # Save the document
    try:
        doc.save(output_path); print(f"\n✅ DOCX saved: {output_path}"); print(f"  Processed: {processed_count}, Skipped: {skipped_count}")
    except Exception as save_e: print(f"❌ Error saving DOCX: {save_e}"); traceback.print_exc()

# ---------------- FORMAT CONVERSION FUNCTIONS ----------------

def convert_pdf_to_docx(pdf_path, docx_path):
    """ Converts PDF to DOCX using pdf2docx. """
    print(f"--- Starting PDF to DOCX Conversion ---")
    print(f"  Input: {pdf_path}")
    print(f"  Output: {docx_path}")
    try:
        from pdf2docx import Converter # Moved import inside try
        # Compatibility shim: some pdf2docx versions call Rect.get_area()
        # but recent PyMuPDF's Rect doesn't provide that method. Add it if missing.
        try:
            if hasattr(fitz, 'Rect') and not hasattr(fitz.Rect, 'get_area'):
                def _rect_get_area(self):
                    try:
                        return float(self.width * self.height)
                    except Exception:
                        # Fallback using bbox coordinates
                        try:
                            return float((self.x1 - self.x0) * (self.y1 - self.y0))
                        except Exception:
                            return 0.0
                setattr(fitz.Rect, 'get_area', _rect_get_area)
        except Exception:
            pass
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)  # type: ignore[arg-type]
        cv.close()
        if not os.path.exists(docx_path) or os.path.getsize(docx_path) < 100: # Basic check
             raise Exception("Output DOCX file seems invalid or empty.")
        print(f"✅ Converted PDF to DOCX successfully.")
    except ImportError:
        print("❌ pdf2docx not installed. Install with: pip install pdf2docx")
        raise
    except Exception as e:
        print(f"❌ PDF to DOCX conversion failed: {e}")
        traceback.print_exc()
        raise


def convert_docx_to_pdf(docx_path, pdf_path):
    """ Converts DOCX to PDF using LibreOffice or comtypes fallback. """
    print(f"--- Starting DOCX to PDF Conversion ---")
    print(f"  Input: {docx_path}")
    print(f"  Output: {pdf_path}")
    output_dir = os.path.dirname(pdf_path)
    os.makedirs(output_dir, exist_ok=True)
    abs_docx_path = os.path.abspath(docx_path)
    abs_pdf_path = os.path.abspath(pdf_path)
    abs_output_dir = os.path.abspath(output_dir)
    libreoffice_paths = []
    libreoffice_exe = None
    system = platform.system()
    libreoffice_used = False # Initialize flag

    try:
        # --- Find LibreOffice ---
        if system == "Windows":
            libreoffice_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
            ]
        elif system == "Darwin": # macOS
             libreoffice_paths.append("/Applications/LibreOffice.app/Contents/MacOS/soffice")
             # Try 'which' as fallback
             try:
                 result = subprocess.run(['which', 'soffice'], capture_output=True, text=True, check=True, timeout=5)
                 soffice_path = result.stdout.strip()
                 if soffice_path: libreoffice_paths.append(soffice_path)
             except Exception: pass # Ignore if 'which' fails
        else: # Linux/Other
            try: # Use 'which'
                result = subprocess.run(['which', 'libreoffice'], capture_output=True, text=True, check=True, timeout=5)
                soffice_path = result.stdout.strip()
                if soffice_path: libreoffice_paths.append(soffice_path)
            except Exception: pass # Ignore if 'which' fails

        # Select the first existing path
        libreoffice_exe = next((path for path in libreoffice_paths if os.path.exists(path)), None)

        # --- Attempt Conversion ---
        if libreoffice_exe:
            print(f"  Using LibreOffice: {libreoffice_exe}")
            # Use subprocess.run with timeout and capture output
            result = subprocess.run([libreoffice_exe, "--headless", "--convert-to", "pdf", "--outdir", abs_output_dir, abs_docx_path],
                                    check=True, capture_output=True, text=True, timeout=180) # Increased timeout
            # Log output only if needed for debugging
            # print(f"  LibreOffice stdout: {result.stdout}")
            # if result.stderr: print(f"  LibreOffice stderr: {result.stderr}")
            libreoffice_used = True
        elif system == "Windows": # comtypes fallback only on Windows
             print("  LibreOffice not found. Attempting MS Word (comtypes) fallback...")
             try:
                 import comtypes.client # type: ignore[import-not-found]
                 word, doc_com = None, None # Renamed doc to avoid conflict
                 try:
                     comtypes.CoInitialize() # type: ignore[attr-defined]
                     word = comtypes.client.CreateObject('Word.Application')
                     word.Visible = False
                     doc_com = word.Documents.Open(abs_docx_path)
                     doc_com.SaveAs(abs_pdf_path, FileFormat=17) # 17 = wdFormatPDF
                     print(f"  Conversion using Word successful.")
                 finally:
                     if doc_com: doc_com.Close(SaveChanges=0)
                     if word: word.Quit()
                     comtypes.CoUninitialize() # type: ignore[attr-defined]
             except ImportError: raise Exception("LibreOffice not found and 'comtypes' not installed.")
             except Exception as com_e: raise Exception(f"comtypes (MS Word) conversion failed: {com_e}")
        else: # Non-Windows without LibreOffice
             raise Exception(f"LibreOffice not found at expected paths ({', '.join(libreoffice_paths)}) and no fallback available on {system}.")

        # --- Post-conversion check/rename (mainly for LibreOffice) ---
        if libreoffice_used: # Check only if LibreOffice was attempted
            expected_output_filename = os.path.basename(abs_docx_path).replace(".docx", ".pdf")
            created_pdf_path = os.path.join(abs_output_dir, expected_output_filename)
            # Check if the file exists at the expected LibreOffice output path
            if os.path.exists(created_pdf_path):
                if created_pdf_path != abs_pdf_path: # Rename if LibreOffice didn't use the exact target path
                    if os.path.exists(abs_pdf_path): os.remove(abs_pdf_path) # Remove if target exists (overwrite)
                    shutil.move(created_pdf_path, abs_pdf_path)
                # Final check after potential move
                if os.path.exists(abs_pdf_path) and os.path.getsize(abs_pdf_path) > 100: # Check size > 100 bytes
                    print(f"✅ Converted DOCX to PDF successfully: {abs_pdf_path}")
                else: raise Exception(f"LibreOffice conversion finished, but final PDF is missing or empty/invalid at {abs_pdf_path} after potential move.")
            # Handle case where LibreOffice might have created the file directly at abs_pdf_path (less common but possible)
            elif os.path.exists(abs_pdf_path) and os.path.getsize(abs_pdf_path) > 100:
                print(f"✅ Converted DOCX to PDF successfully (likely created directly): {abs_pdf_path}")
            else: # If neither path has a valid file
                raise Exception(f"LibreOffice conversion finished, but no valid output PDF found at expected locations ({created_pdf_path} or {abs_pdf_path}). Check LibreOffice output/errors.")

    # --- Error Handling ---
    except FileNotFoundError as fnf_e:
        print(f"❌ DOCX to PDF conversion failed: {fnf_e}")
        print("ℹ️ Ensure LibreOffice is installed and accessible in PATH or standard locations.")
        raise
    except subprocess.CalledProcessError as cpe:
        print(f"❌ DOCX to PDF conversion failed (LibreOffice Error): {cpe}")
        # Decode stderr/stdout safely, handling potential None or non-byte values
        stderr = cpe.stderr.decode('utf-8', errors='ignore') if isinstance(cpe.stderr, bytes) else str(cpe.stderr or '')
        stdout = cpe.stdout.decode('utf-8', errors='ignore') if isinstance(cpe.stdout, bytes) else str(cpe.stdout or '')
        print(f"  LibreOffice stdout: {stdout}")
        print(f"  LibreOffice stderr: {stderr}")
        raise
    except subprocess.TimeoutExpired:
         print(f"❌ DOCX to PDF conversion timed out after 180 seconds.")
         raise
    except Exception as e:
        print(f"❌ DOCX to PDF conversion failed with an unexpected error: {e}")
        traceback.print_exc()
        if "comtypes" not in str(e) and "LibreOffice" not in str(e): # Avoid redundant message
            print("ℹ️ Ensure LibreOffice is installed OR (on Windows) MS Word + comtypes library are available.")
        raise